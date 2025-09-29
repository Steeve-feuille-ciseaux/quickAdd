from django.shortcuts import render, redirect, get_object_or_404
from django.utils import timezone
from .forms import EtudiantForm
from .models import Etudiant
from datetime import date
from collections import Counter
from django.http import HttpResponse
from docx import Document
from io import BytesIO
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def ajouter_etudiant(request):
    if request.method == 'POST':
        form = EtudiantForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('liste_etudiants')
    else:
        form = EtudiantForm()
    return render(request, 'club/ajouter_etudiant.html', {'form': form})

def liste_etudiants(request):
    etudiants_qs = Etudiant.objects.all()
    total_inscrits = etudiants_qs.count()
    niveaux = etudiants_qs.values_list('niveau', flat=True)
    compteur = Counter(niveaux)

    # Garantir toutes les clés 6,5,4,3
    niveaux_fixes = ['6', '5', '4', '3']
    inscrits_par_niveau = {niveau: compteur.get(niveau, 0) for niveau in niveaux_fixes}

    # ✅ Date d’ouverture des inscriptions
    date_ouverture = date(2025, 9, 25)

    # ✅ Préparer liste enrichie avec délai
    etudiants = []
    for etudiant in etudiants_qs:
        # Calcul du délai
        delta = (etudiant.date_inscription - date_ouverture).days
        etudiants.append({
            'instance': etudiant,
            'delai': f"(J+{delta})" if delta >= 0 else f"(J-{abs(delta)})",  # Si inscrit avant ouverture
        })

    context = {
        'etudiants': etudiants,
        'total_inscrits': total_inscrits,
        'inscrits_par_niveau': inscrits_par_niveau,
    }
    return render(request, 'club/liste_etudiants.html', context)

def appel(request):
    etudiants = Etudiant.objects.all().order_by('nom')
    return render(request, 'club/appel.html', {'etudiants': etudiants})

def valider_presence(request, etudiant_id):
    etudiant = get_object_or_404(Etudiant, id=etudiant_id)
    etudiant.date_presence = timezone.now().date()
    etudiant.cours_suivi += 1
    etudiant.save()
    return redirect('appel')

def export_word(request):
    etudiants = Etudiant.objects.all()

    document = Document()
    heading = document.add_heading('Club informatique', level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Ajouter tableau résumé
    total_inscrits = etudiants.count()
    niveaux = etudiants.values_list('niveau', flat=True)
    from collections import Counter
    inscrits_par_niveau = dict(Counter(niveaux))

    # Résumé
    # Supposons que tu as tous les niveaux dans une liste (ordre voulu)
    niveaux_possibles = ['6', '5', '4', '3']

    # Création de la table résumé avec tous les niveaux
    table_resume = document.add_table(rows=2, cols=1 + len(niveaux_possibles))
    hdr_cells = table_resume.rows[0].cells
    hdr_cells[0].text = 'Total '

    for i, niveau in enumerate(niveaux_possibles, start=1):
        hdr_cells[i].text = f'{niveau}ième'

    row_cells = table_resume.rows[1].cells
    row_cells[0].text = str(total_inscrits)

    for i, niveau in enumerate(niveaux_possibles, start=1):
        count = inscrits_par_niveau.get(niveau, 0)
        row_cells[i].text = str(count)

    document.add_paragraph()  # saut de ligne

    # Tableau détails avec colonnes : Prénom, Nom, Niveau+Classe, Thème
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'  # style simple avec bordures

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Prénom'
    hdr_cells[1].text = 'Nom'
    hdr_cells[2].text = 'Classe'
    hdr_cells[3].text = 'Thème'

    # Définir largeurs de colonnes (optionnel, pour mieux répartir l’espace)
    widths = [Inches(1.2), Inches(1.5), Inches(1.8), Inches(3.0)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    for etudiant in etudiants:
        row_cells = table.add_row().cells
        row_cells[0].text = etudiant.prenom
        row_cells[1].text = etudiant.nom
        # Niveau + Classe en majuscule comme demandé
        row_cells[2].text = f"{etudiant.niveau} {etudiant.classe.upper()}"
        row_cells[3].text = etudiant.get_theme_display()
        # row_cells[4].text = etudiant.date_inscription.strftime('%d/%m/%Y')

    # (Optionnel) ajuster la taille de la police pour que ce soit plus lisible
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)

    # Préparer le document à retourner en téléchargement
    f = BytesIO()
    document.save(f)
    f.seek(0)

    response = HttpResponse(
        f.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=liste_etudiants.docx'
    return response
